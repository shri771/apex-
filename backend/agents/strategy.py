import json
import logging
import os
from collections import defaultdict
from datetime import date, datetime, timedelta

from .base import BaseAgent

logger = logging.getLogger(__name__)

SYSTEM_PROMPT = "You are a senior market strategy analyst. Respond only with valid JSON."

BRIEF_DIR = os.path.expanduser("~/storage/downloads/market_briefs")

SECTIONS = [
    "Executive Summary",
    "Top Threats",
    "Top Opportunities",
    "Recommended Actions",
    "Metrics",
]

USER_PROMPT_TEMPLATE = """You have received market intelligence data from the past 7 days. Write a weekly strategy brief.

Return a JSON object with exactly these top-level keys:
{{
  "Executive Summary": "2-5 sentences summarising the market situation",
  "Top Threats": "bullet-point prose of top 3 threats",
  "Top Opportunities": "bullet-point prose of top 3 opportunities",
  "Recommended Actions": "prose summary of actions",
  "Metrics": "key numbers and metrics observed",
  "recommended_actions": [
    {{
      "action": "specific action title",
      "detail": "why this matters and how to execute",
      "impact": "high|medium|low",
      "confidence": 85
    }}
  ],
  "top_recommendation": "single most important action as a short sentence"
}}

Rules:
- recommended_actions must be a JSON array of 3-5 items
- confidence is integer 0-100
- impact is one of: high, medium, low
- top_recommendation is a single concise sentence (max 15 words)

DATA:
{data_block}"""


class StrategyAgent(BaseAgent):
    agent_name = "strategy"

    async def fetch_sources(self) -> list[dict]:
        from backend.db.database import SessionLocal
        from backend.db.models import Insight

        since = datetime.utcnow() - timedelta(days=7)
        week_start = (date.today() - timedelta(days=date.today().weekday())).isoformat()

        db = SessionLocal()
        try:
            rows = (
                db.query(Insight)
                .filter(Insight.agent != "strategy", Insight.created_at >= since)
                .order_by(Insight.agent, Insight.category)
                .all()
            )
            grouped: dict[str, list[dict]] = defaultdict(list)
            for row in rows:
                grouped[row.agent].append({
                    "summary": row.summary,
                    "category": row.category,
                    "severity": row.severity,
                    "score": row.score,
                    "source": row.source,
                })
        finally:
            db.close()

        return [{"grouped": dict(grouped), "week_start": week_start}]

    async def analyse(self, item: dict) -> dict:
        grouped = item.get("grouped", {})
        week_start = item.get("week_start", date.today().isoformat())

        data_block = self._build_data_block(grouped)

        user_msg = USER_PROMPT_TEMPLATE.format(data_block=data_block)
        model = self._get_setting("ollama_model", "phi3:mini")
        result = await self.ollama_client.generate(
            model=model,
            system=SYSTEM_PROMPT,
            user=user_msg,
            timeout=180,
        )

        sections = {}
        for section in SECTIONS:
            sections[section] = result.get(section, "No data available.")

        # Extract structured recommended_actions with fallback
        raw_actions = result.get("recommended_actions", [])
        if not isinstance(raw_actions, list):
            raw_actions = []

        # Normalise each action
        recommended_actions = []
        for a in raw_actions[:5]:
            if isinstance(a, dict):
                recommended_actions.append({
                    "action": str(a.get("action", "")),
                    "detail": str(a.get("detail", "")),
                    "impact": str(a.get("impact", "medium")).lower(),
                    "confidence": max(0, min(100, int(a.get("confidence", 70)))),
                })

        top_recommendation = str(result.get("top_recommendation", "Review market intelligence data."))

        self._write_docx(sections, week_start)
        self._insert_brief_record(sections, week_start)

        # Derive category from data: more threats → threat, more opportunities → opportunity
        threat_count = sum(
            1 for agent_items in grouped.values()
            for f in agent_items if f.get("category") == "threat"
        )
        opp_count = sum(
            1 for agent_items in grouped.values()
            for f in agent_items if f.get("category") == "opportunity"
        )
        derived_category = "threat" if threat_count > opp_count else "opportunity"

        executive_summary = sections.get("Executive Summary", "Weekly brief generated.")
        return {
            "summary": executive_summary[:500],
            "category": derived_category,
            "severity": "low",
            "score": 0.5,
            "key_points": list(sections.keys()),
            "source": f"brief_{week_start}.docx",
            "raw_text": data_block[:2000],
            "metadata": json.dumps({
                "week_start": week_start,
                "sections": sections,
                "recommended_actions": recommended_actions,
                "top_recommendation": top_recommendation,
            }),
        }

    def _build_data_block(self, grouped: dict) -> str:
        lines = []
        for agent_name, findings in grouped.items():
            lines.append(f"\n### {agent_name.upper()} FINDINGS ({len(findings)} items)")
            by_cat: dict[str, list] = defaultdict(list)
            for f in findings:
                by_cat[f.get("category", "neutral")].append(f)
            for cat, items in by_cat.items():
                lines.append(f"\n{cat.upper()}:")
                for item in items[:5]:
                    sev = item.get("severity", "low")
                    lines.append(f"  - [{sev}] {item.get('summary', '')}")
        return "\n".join(lines)

    def _write_docx(self, sections: dict, week_start: str):
        try:
            from docx import Document
            from docx.shared import Pt

            os.makedirs(BRIEF_DIR, exist_ok=True)
            doc = Document()
            doc.add_heading(f"Market Intelligence Brief — Week of {week_start}", level=0)

            for section_title in SECTIONS:
                content = sections.get(section_title, "")
                doc.add_heading(section_title, level=1)
                para = doc.add_paragraph(content)
                if para.runs:
                    para.runs[0].font.size = Pt(11)

            file_path = os.path.join(BRIEF_DIR, f"brief_{week_start}.docx")
            doc.save(file_path)
            logger.info("strategy: brief saved to %s", file_path)
        except Exception as exc:
            logger.error("strategy: failed to write .docx: %s", exc)

    def _insert_brief_record(self, sections: dict, week_start: str):
        from backend.db.database import SessionLocal
        from backend.db.models import Brief

        file_path = os.path.join(BRIEF_DIR, f"brief_{week_start}.docx")
        summary = sections.get("Executive Summary", "")[:500]

        db = SessionLocal()
        try:
            existing = db.query(Brief).filter(Brief.week_start == week_start).first()
            if existing:
                existing.file_path = file_path
                existing.summary = summary
            else:
                db.add(Brief(week_start=week_start, file_path=file_path, summary=summary))
            db.commit()
        except Exception as exc:
            logger.error("strategy: failed to insert brief record: %s", exc)
            db.rollback()
        finally:
            db.close()
